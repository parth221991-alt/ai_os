from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand tokens ─────────────────────────────────────────────────────────────
EMERALD     = RGBColor(0x10, 0xB9, 0x81)   # #10B981
DARK        = RGBColor(0x0F, 0x17, 0x2A)   # near-black
MID         = RGBColor(0x47, 0x55, 0x69)   # #475569 slate
LIGHT_BG    = RGBColor(0xF0, 0xFD, 0xF4)   # #F0FDF4 emerald-50
CODE_BG     = RGBColor(0xF1, 0xF5, 0xF9)   # #F1F5F9 slate-100
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
EMERALD_DARK= RGBColor(0x06, 0x5F, 0x46)   # #065F46


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)


def add_para_border_bottom(para, color='10B981', size=12):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(size))
    bot.set(qn('w:space'), '6')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def set_para_shading(para, rgb: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    pPr.append(shd)


def add_run(para, text, bold=False, italic=False,
            size=11, color=None, font='Calibri'):
    run      = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size    = Pt(size)
    run.font.name    = font
    run.font.color.rgb = color or DARK
    return run


def add_heading(doc, text, level=1, bottom_border=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    para.paragraph_format.space_after  = Pt(6)
    if level == 1:
        add_run(para, text, bold=True, size=20, color=DARK, font='Calibri')
        if bottom_border:
            add_para_border_bottom(para)
    elif level == 2:
        add_run(para, '▌ ', bold=True, size=14, color=EMERALD, font='Calibri')
        add_run(para, text, bold=True, size=14, color=DARK, font='Calibri')
    elif level == 3:
        add_run(para, text, bold=True, size=12, color=EMERALD_DARK, font='Calibri')
    return para


def add_body(doc, text, color=None, italic=False, size=11):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(6)
    add_run(para, text, italic=italic, size=size, color=color or MID)
    return para


def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.left_indent  = Inches(0.3 + level * 0.25)
    add_run(para, text, size=11, color=MID)
    return para


def add_code_block(doc, code_text):
    for line in code_text.split('\n'):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.left_indent  = Inches(0.3)
        para.paragraph_format.right_indent = Inches(0.3)
        set_para_shading(para, CODE_BG)
        run = para.add_run(line if line else ' ')
        run.font.name  = 'Consolas'
        run.font.size  = Pt(10)
        run.font.color.rgb = DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_step_badge(doc, number, title, subtitle=''):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.allow_autofit = False

    # Badge cell
    badge_cell = tbl.cell(0, 0)
    badge_cell.width = Inches(0.7)
    set_cell_bg(badge_cell, EMERALD)
    badge_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    bp = badge_cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.paragraph_format.space_before = Pt(6)
    bp.paragraph_format.space_after  = Pt(6)
    r = bp.add_run(f'0{number}')
    r.font.name  = 'Calibri'
    r.font.size  = Pt(22)
    r.font.bold  = True
    r.font.color.rgb = WHITE

    # Title cell
    title_cell = tbl.cell(0, 1)
    set_cell_bg(title_cell, LIGHT_BG)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tp = title_cell.paragraphs[0]
    tp.paragraph_format.left_indent  = Inches(0.15)
    tp.paragraph_format.space_before = Pt(6)
    tp.paragraph_format.space_after  = Pt(2)
    t1 = tp.add_run(title)
    t1.font.name  = 'Calibri'
    t1.font.size  = Pt(15)
    t1.font.bold  = True
    t1.font.color.rgb = DARK
    if subtitle:
        tp.add_run('\n')
        t2 = tp.add_run(subtitle)
        t2.font.name   = 'Calibri'
        t2.font.size   = Pt(10)
        t2.font.italic = True
        t2.font.color.rgb = MID
        tp.paragraph_format.space_after = Pt(6)

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    return tbl


def add_divider(doc, color='10B981'):
    para = doc.add_paragraph()
    add_para_border_bottom(para, color=color, size=6)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(8)


def set_page_margins(doc, top=1.0, bottom=1.0, left=1.15, right=1.15):
    section = doc.sections[0]
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)


def add_header_footer(doc):
    # Header
    section = doc.sections[0]
    header  = section.header
    hp      = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = hp.add_run('AI_SNIPP')
    r1.font.name  = 'Calibri'
    r1.font.size  = Pt(10)
    r1.font.bold  = True
    r1.font.color.rgb = EMERALD
    r2 = hp.add_run('  ·  @ai_snipp  ·  Claude OS Setup Guide')
    r2.font.name  = 'Calibri'
    r2.font.size  = Pt(9)
    r2.font.color.rgb = MID

    # Footer
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f1 = fp.add_run('© 2026 AI_SNIPP  ·  Follow ')
    f1.font.name  = 'Calibri'
    f1.font.size  = Pt(9)
    f1.font.color.rgb = MID
    f2 = fp.add_run('@ai_snipp')
    f2.font.name  = 'Calibri'
    f2.font.size  = Pt(9)
    f2.font.bold  = True
    f2.font.color.rgb = EMERALD
    f3 = fp.add_run('  for daily AI workflows  ·  Not for redistribution')
    f3.font.name  = 'Calibri'
    f3.font.size  = Pt(9)
    f3.font.color.rgb = MID


# ── Document ──────────────────────────────────────────────────────────────────

doc = Document()
set_page_margins(doc)
add_header_footer(doc)

# ── Cover banner ──────────────────────────────────────────────────────────────
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = tbl.cell(0, 0)
set_cell_bg(cell, DARK)
cp = cell.paragraphs[0]
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.paragraph_format.space_before = Pt(16)
cp.paragraph_format.space_after  = Pt(4)
r = cp.add_run('AI_SNIPP')
r.font.name  = 'Calibri'
r.font.size  = Pt(11)
r.font.bold  = True
r.font.color.rgb = EMERALD
cp2 = cell.add_paragraph()
cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = cp2.add_run('Build Your Claude OS in 3 Steps')
r2.font.name  = 'Calibri'
r2.font.size  = Pt(26)
r2.font.bold  = True
r2.font.color.rgb = WHITE
cp3 = cell.add_paragraph()
cp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp3.paragraph_format.space_after = Pt(16)
r3 = cp3.add_run('A personal operating system that gives Claude permanent memory of your work')
r3.font.name   = 'Calibri'
r3.font.size   = Pt(12)
r3.font.italic = True
r3.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)   # slate-400

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Intro ──────────────────────────────────────────────────────────────────────
add_body(doc,
    'Most Claude users re-explain their entire context at the start of every session. '
    'This guide shows you how to stop doing that — permanently.',
    size=12)

add_body(doc,
    'The system takes about 20 minutes to set up. After that, Claude already knows who you are, '
    'what you\'re building, and how to think — every time you open it.',
    size=12)

# System logic callout
tbl2 = doc.add_table(rows=1, cols=1)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
c2 = tbl2.cell(0, 0)
set_cell_bg(c2, LIGHT_BG)
pp = c2.paragraphs[0]
pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
pp.paragraph_format.space_before = Pt(10)
pp.paragraph_format.space_after  = Pt(10)
rr = pp.add_run('Structure  →  Brief  →  Activate')
rr.font.name  = 'Calibri'
rr.font.size  = Pt(15)
rr.font.bold  = True
rr.font.color.rgb = EMERALD_DARK
pp.add_run('\n')
rr2 = pp.add_run('One folder system. One file. One command. That\'s the entire system.')
rr2.font.name   = 'Calibri'
rr2.font.size   = Pt(10)
rr2.font.italic = True
rr2.font.color.rgb = MID

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_divider(doc)

# ── Step 1 ────────────────────────────────────────────────────────────────────
add_step_badge(doc, 1, 'CREATE THE FOLDER STRUCTURE', 'Time: ~20 minutes')

add_body(doc,
    'Create a root folder on your machine — name it anything you\'ll remember. '
    'Inside it, create exactly 9 numbered sub-folders:')

add_code_block(doc,
"""YourWorkspace/
├── 01_memory/      ← Notes and decisions you want Claude to retain
├── 02_skills/      ← Reusable task definitions for recurring work
├── 03_prompts/     ← Tested prompt templates you want to reuse
├── 04_projects/    ← Your actual project folders (or symlinks)
├── 05_content/     ← Drafts, scripts, research outputs
├── 06_agents/      ← Agent configs for multi-step automations
├── 07_templates/   ← Starter templates for common project types
├── 08_mcp/         ← MCP server configs (Claude's external tools)
└── 09_docs/        ← Architecture notes and auditable decisions""")

add_heading(doc, 'Why the Numbers?', level=3)
add_body(doc,
    'Claude Code reads folders in alphabetical order. Numbering forces the priority — '
    '01_memory loads before 09_docs. Your most important context is always read first.')

add_heading(doc, 'Where to Start', level=3)
add_bullet(doc, 'You don\'t need to fill all 9 folders immediately')
add_bullet(doc, 'Start with 01_memory, 03_prompts, and 04_projects')
add_bullet(doc, 'The rest populates naturally as you work')

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_divider(doc)

# ── Step 2 ────────────────────────────────────────────────────────────────────
add_step_badge(doc, 2, 'WRITE YOUR CLAUDE.MD', 'Time: ~30 minutes — most important step')

add_body(doc,
    'CLAUDE.md is a single markdown file saved directly in your root workspace folder '
    '(not inside any sub-folder). This is the file that changes everything.')

add_code_block(doc, 'YourWorkspace/CLAUDE.md   ← lives here, not inside any sub-folder')

add_body(doc,
    'Claude Code reads this file automatically every time you open a session in this folder. '
    'It is your permanent brief — you write it once and stop re-explaining forever.')

add_heading(doc, 'What to Include in CLAUDE.md', level=2)

add_heading(doc, 'Section 1 — Who You Are', level=3)
add_body(doc,
    'One paragraph. Your role, domain, experience level. Be specific — "freelance developer" '
    'is less useful than "freelance React developer building fintech dashboards."')

add_heading(doc, 'Section 2 — What You\'re Building', level=3)
add_body(doc,
    'List your active projects, one sentence each: what it is, what stack, current status. '
    'Claude uses this to give contextually relevant answers without needing to ask.')

add_heading(doc, 'Section 3 — How to Think', level=3)
add_body(doc,
    'This section is the most powerful — and the most skipped. Tell Claude your tradeoffs, '
    'your non-negotiables, and how you like feedback. Examples:')
add_bullet(doc, '"Correctness first. Simplicity second. Performance only when measured."')
add_bullet(doc, '"If I ask for a fix, fix only the broken thing — don\'t refactor."')
add_bullet(doc, '"Tell me when my approach is wrong before writing any code."')
add_bullet(doc, '"No placeholder comments. No TODO stubs. Complete implementations only."')

add_heading(doc, 'Section 4 — Technical Standards  (optional)', level=3)
add_body(doc,
    'Your preferred stack, naming conventions, file structure rules, testing approach. '
    'Anything that applies to every task — list it here so Claude never asks.')

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_heading(doc, 'CLAUDE.md Template — Copy and Fill In', level=2)

add_code_block(doc,
"""# My Claude Workspace

## Who I Am
[Your role in 2–3 sentences. Be specific about domain and experience level.]

## Active Projects
- **Project A** — [What it is. Stack. Current status.]
- **Project B** — [What it is. Stack. Current status.]

## How to Think
- Correctness first. Then simplicity. Then performance.
- If I ask for a fix, fix only the broken thing — don't refactor.
- Tell me when my approach is wrong before writing code.
- No placeholder comments. No TODO stubs. Complete implementations only.
- When in doubt, ask one clarifying question — don't guess.

## Technical Standards
- Language / framework: [your stack]
- Linting / formatting: [your rules]
- Testing approach: [your standard]
- Non-negotiables: [anything that applies to every task]""")

add_body(doc,
    'The more specific you are, the less you\'ll repeat yourself. A vague CLAUDE.md '
    'gives vague sessions. A specific one gives a working partner that already knows your context.',
    italic=True, color=MID)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_divider(doc)

# ── Step 3 ────────────────────────────────────────────────────────────────────
add_step_badge(doc, 3, 'ACTIVATE WITH CLAUDE CODE', 'Time: ~5 minutes')

add_body(doc,
    'Once your folder structure and CLAUDE.md exist, open Claude Code inside your workspace folder. '
    'Claude Code is Anthropic\'s official CLI — it reads CLAUDE.md automatically on startup.')

add_heading(doc, 'Install Claude Code', level=3)
add_code_block(doc, 'npm install -g @anthropic-ai/claude-code')

add_heading(doc, 'Open in Your Workspace', level=3)
add_code_block(doc, 'cd /path/to/YourWorkspace\nclaude')

add_heading(doc, 'What Happens on First Launch', level=3)
add_body(doc,
    'Claude reads your CLAUDE.md. The first response references your context directly — '
    'your projects, your standards, your preferences. You don\'t explain anything. It already knows.')

add_heading(doc, 'Requirements', level=3)
add_bullet(doc, 'Node.js installed (for npm)')
add_bullet(doc, 'Claude Pro subscription OR Anthropic API key')
add_bullet(doc, 'The CLI itself is free — you pay only for what you use via API')

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_divider(doc)

# ── What Happens After ────────────────────────────────────────────────────────
add_heading(doc, 'What Changes After Activation', level=1, bottom_border=True)

add_body(doc,
    'Every Claude session inside your workspace folder starts with full context — automatically. '
    'No re-introduction. No setup. No repeated instructions.')

pairs = [
    ('Before',                                  'After'),
    ('Re-explain your project every session',   'Claude already knows your active projects'),
    ('Repeat your coding standards each time',  'Standards are applied from CLAUDE.md'),
    ('Describe your role from scratch',         'Claude knows who it\'s working with'),
    ('Paste the same system prompt repeatedly', 'System prompt is permanent — in the file'),
    ('Start fresh after a long break',          'CLAUDE.md is always there — nothing resets'),
]

tbl3 = doc.add_table(rows=len(pairs), cols=2)
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl3.style     = 'Table Grid'

for i, (before, after) in enumerate(pairs):
    cell_b = tbl3.cell(i, 0)
    cell_a = tbl3.cell(i, 1)
    if i == 0:
        set_cell_bg(cell_b, DARK)
        set_cell_bg(cell_a, DARK)
        pb = cell_b.paragraphs[0]
        pa = cell_a.paragraphs[0]
        pb.paragraph_format.space_before = Pt(4)
        pb.paragraph_format.space_after  = Pt(4)
        pa.paragraph_format.space_before = Pt(4)
        pa.paragraph_format.space_after  = Pt(4)
        rb = pb.add_run(before)
        ra = pa.add_run(after)
        rb.font.name = ra.font.name = 'Calibri'
        rb.font.size = ra.font.size = Pt(11)
        rb.font.bold = ra.font.bold = True
        rb.font.color.rgb = ra.font.color.rgb = WHITE
    else:
        set_cell_bg(cell_b, RGBColor(0xFF, 0xF1, 0xF2))  # soft red
        set_cell_bg(cell_a, LIGHT_BG)
        pb = cell_b.paragraphs[0]
        pa = cell_a.paragraphs[0]
        pb.paragraph_format.space_before = Pt(4)
        pb.paragraph_format.space_after  = Pt(4)
        pa.paragraph_format.space_before = Pt(4)
        pa.paragraph_format.space_after  = Pt(4)
        rb = pb.add_run('✗  ' + before)
        ra = pa.add_run('✓  ' + after)
        rb.font.name = ra.font.name = 'Calibri'
        rb.font.size = ra.font.size = Pt(10)
        rb.font.color.rgb = RGBColor(0x9B, 0x1C, 0x1C)   # red-800
        ra.font.color.rgb = EMERALD_DARK

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ── Pro Tips ──────────────────────────────────────────────────────────────────
add_heading(doc, 'Pro Tips', level=1, bottom_border=True)

tips = [
    ('Start specific, not perfect',
     'A rough but honest CLAUDE.md beats a polished generic one. '
     'Write what\'s actually true about your work today. Refine it as you go.'),
    ('Use 01_memory as a running log',
     'Drop key decisions, learnings, and constraints into 01_memory as plain text files. '
     'Claude reads everything in there. It becomes your project knowledge base over time.'),
    ('Version control the whole workspace',
     'Run git init in your workspace root. Commit CLAUDE.md changes with a note on what you changed '
     'and why. This makes your context history auditable.'),
    ('One workspace per domain, not per project',
     'Don\'t create a separate workspace for every project. One workspace with all projects listed '
     'in CLAUDE.md is cleaner and prevents context fragmentation.'),
]

for title, body in tips:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run('→  ')
    r1.font.name = 'Calibri'; r1.font.size = Pt(12); r1.font.bold = True; r1.font.color.rgb = EMERALD
    r2 = p.add_run(title)
    r2.font.name = 'Calibri'; r2.font.size = Pt(12); r2.font.bold = True; r2.font.color.rgb = DARK
    add_body(doc, body, size=10)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_divider(doc)

# ── CTA Banner ────────────────────────────────────────────────────────────────
tbl4 = doc.add_table(rows=1, cols=1)
c4   = tbl4.cell(0, 0)
set_cell_bg(c4, DARK)
fp4  = c4.paragraphs[0]
fp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp4.paragraph_format.space_before = Pt(14)
fp4.paragraph_format.space_after  = Pt(4)
rc1 = fp4.add_run('Want more AI workflows like this?')
rc1.font.name = 'Calibri'; rc1.font.size = Pt(13); rc1.font.bold = True; rc1.font.color.rgb = WHITE
sp4 = c4.add_paragraph()
sp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp4.paragraph_format.space_after = Pt(14)
rc2 = sp4.add_run('Follow  ')
rc2.font.name = 'Calibri'; rc2.font.size = Pt(13); rc2.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
rc3 = sp4.add_run('@ai_snipp')
rc3.font.name = 'Calibri'; rc3.font.size = Pt(16); rc3.font.bold = True; rc3.font.color.rgb = EMERALD
rc4 = sp4.add_run('  on Instagram · Daily AI workflows, tools, and systems')
rc4.font.name = 'Calibri'; rc4.font.size = Pt(11); rc4.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ── Save ──────────────────────────────────────────────────────────────────────
out = r'D:\AI_OS\05_content\AI_SNIPP\REEL_021_AI_OS_guide.docx'
doc.save(out)
print(f'Saved: {out}')
